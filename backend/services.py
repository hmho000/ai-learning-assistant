import fitz  # PyMuPDF
import re
import json
import requests
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from .models import Chapter, Quiz, Question
from sqlmodel import Session

# === 常量 ===
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# === PDF 解析服务 ===

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    简单提取 PDF 全文文本
    """
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def parse_chapters_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """
    尝试从 PDF 目录提取章节
    返回: [{"title": "第1章...", "content": "..."}]
    """
    doc = fitz.open(pdf_path)
    toc = doc.get_toc()
    
    chapters = []
    
    # 如果没有目录，这就比较麻烦，暂时只做一个简单的全书作为一个章节
    if not toc:
        full_text = extract_text_from_pdf(pdf_path)
        chapters.append({
            "title": "全书内容",
            "index": 1,
            "content": full_text
        })
        return chapters

    # 简单的目录解析逻辑
    # 寻找一级目录
    level1_nodes = [item for item in toc if item[0] == 1]
    
    for i, node in enumerate(level1_nodes):
        title = node[1]
        start_page = node[2]
        
        # 确定结束页码
        if i < len(level1_nodes) - 1:
            end_page = level1_nodes[i+1][2]
        else:
            end_page = doc.page_count
            
        # 提取该范围的文本
        chapter_text = ""
        # fitz 页码从 0 开始，toc 从 1 开始
        for p in range(start_page - 1, end_page - 1):
            if p < doc.page_count:
                chapter_text += doc[p].get_text()
        
        chapters.append({
            "title": title,
            "index": i + 1,
            "content": chapter_text
        })
        
    return chapters


def generate_quiz_for_chapter(chapter_text: str, chapter_title: str, 
                              num_mc: int = 5, 
                              num_multi: int = 0,
                              num_tf: int = 0,
                              num_fb: int = 5,
                              num_short: int = 0,
                              num_code: int = 0,
                              difficulty: str = "medium") -> Dict[str, Any]:
    """
    调用 DeepSeek 生成题目
    """
    # 优先从环境变量获取，如果没有则报错
    # 确保加载了环境变量
    from dotenv import load_dotenv
    env_path = Path(__file__).parent.parent / ".env"
    load_dotenv(dotenv_path=env_path, override=True)
    
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key or api_key.strip() == "sk-your_api_key_here":
        print(f"[ERROR] Invalid API Key. Tried loading from: {env_path.resolve()}")
        raise ValueError("未配置有效的 DEEPSEEK_API_KEY，请检查 .env 文件。")
    
    prompt = f"""
你是一名专业的教育测评专家。
请阅读以下章节内容（标题：{chapter_title}），并根据难度【{difficulty}】生成以下题目：
- 单选题: {num_mc} 道
- 多选题: {num_multi} 道
- 判断题: {num_tf} 道
- 填空题: {num_fb} 道
- 简答题: {num_short} 道
- 代码题: {num_code} 道

要求：
1. 题目必须基于提供的文本。
2. 输出必须是合法的 JSON 格式。
3. 不要包含 markdown 标记。
4. 多选题的 answer 应为包含正确选项字母的数组，如 ["A", "C"]。
5. 判断题的 answer 应为 "True" 或 "False"。
6. 简答题需提供参考答案 (answer) 和评分关键词 (keywords)，且 explanation 必须包含原文引用 (Source Quote)。
7. 代码题需提供题目描述 (question)、参考代码 (answer) 和测试用例说明 (explanation)。

JSON 结构示例：
{{
  "quiz_title": "...",
  "quiz_description": "...",
  "multiple_choice": [
    {{ "question": "...", "options": ["A...", "B...", "C...", "D..."], "answer": "A", "explanation": "..." }}
  ],
  "multi_select": [
    {{ "question": "...", "options": ["A...", "B...", "C...", "D..."], "answer": ["A", "C"], "explanation": "..." }}
  ],
  "true_false": [
    {{ "question": "...", "answer": "True", "explanation": "..." }}
  ],
  "fill_in_blank": [
    {{ "question": "...", "answer": "...", "explanation": "..." }}
  ],
  "short_answer": [
    {{ "question": "...", "answer": "...", "keywords": ["...", "..."], "explanation": "..." }}
  ],
  "coding": [
    {{ "question": "...", "answer": "def func():...", "explanation": "..." }}
  ]
}}

【章节内容开始】
{chapter_text[:8000]} 
【章节内容结束】
(注：内容已截断，仅供参考)
"""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是一个辅助出题的 AI 助手。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.3,
        "stream": False,
        "max_tokens": 8192
    }
    
    try:
        print(f"Sending request to DeepSeek API for chapter: {chapter_title}...")
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=120) # 增加超时时间
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        
        # 清理 markdown 标记
        if content.startswith("```"):
            content = content.strip("`")
            if content.startswith("json"):
                content = content[4:]
        
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            print(f"[ERROR] JSON Parse Error: {e}")
            print(f"[DEBUG] Raw Content (First 500 chars): {content[:500]}")
            print(f"[DEBUG] Raw Content (Last 500 chars): {content[-500:]}")
            raise RuntimeError(f"题目生成返回了无效的 JSON 格式: {e}")

    except requests.exceptions.Timeout:
        raise RuntimeError("DeepSeek API 请求超时，请稍后重试。")
    except requests.exceptions.RequestException as e:
        error_msg = f"DeepSeek API 调用失败: {e}"
        if e.response is not None:
             error_msg += f" (Status: {e.response.status_code})"
        raise RuntimeError(error_msg)
    except Exception as e:
        raise RuntimeError(f"题目生成发生错误: {e}")

def save_quiz_to_db(session: Session, chapter_id: int, quiz_data: Dict[str, Any]):
    """
    将生成的 JSON 数据保存到数据库
    """
    quiz = Quiz(
        chapter_id=chapter_id,
        title=quiz_data.get("quiz_title", "自动生成的练习"),
        description=quiz_data.get("quiz_description", "AI 智能生成")
    )
    session.add(quiz)
    session.commit()
    session.refresh(quiz)
    
    # 保存题目的辅助函数
    def save_q(q_list, q_type):
        for q in q_list:
            # 处理选项 JSON
            options = q.get("options")
            options_json = json.dumps(options, ensure_ascii=False) if options else None
            
            # 处理多选题答案（列表 -> JSON 字符串）
            answer = q["answer"]
            if isinstance(answer, list):
                answer = json.dumps(answer, ensure_ascii=False)
            
            # 处理简答题的额外数据（关键词）...
            # 暂时将其放入 options_json 或 explanation 中？
            # 目前我们坚持使用标准字段。
            
            question = Question(
                quiz_id=quiz.id,
                type=q_type,
                stem=q["question"],
                options_json=options_json,
                answer=str(answer),
                explanation=q.get("explanation")
            )
            session.add(question)

    save_q(quiz_data.get("multiple_choice", []), "multiple_choice")
    save_q(quiz_data.get("multi_select", []), "multi_select")
    save_q(quiz_data.get("true_false", []), "true_false")
    save_q(quiz_data.get("fill_in_blank", []), "fill_in_blank")
    save_q(quiz_data.get("short_answer", []), "short_answer")
    save_q(quiz_data.get("coding", []), "coding")
        
    session.commit()
    return quiz

def grade_short_answer(question_text: str, reference_answer: str, student_answer: str) -> Dict[str, Any]:
    """
    AI 评分简答题
    """
    from dotenv import load_dotenv
    env_path = Path(__file__).parent.parent / ".env"
    load_dotenv(dotenv_path=env_path, override=True)
    api_key = os.getenv("DEEPSEEK_API_KEY")
    prompt = f"""
请对学生的简答题答案进行评分。
题目：{question_text}
参考答案：{reference_answer}
学生答案：{student_answer}

要求：
1. 给出 0-10 分的评分。
2. 评语应简洁，指出学生答案的准确性及遗漏点。

输出 JSON 格式：{{ "score": 8, "feedback": "..." }}
"""
    # ... call API (simplified) ...
    # 为了简洁，重用请求逻辑或假设有一个辅助函数会更好。
    # 我现在将内联实现。
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    try:
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=30)
        content = resp.json()["choices"][0]["message"]["content"]
        if content.startswith("```"): content = content.strip("`").replace("json", "")
        return json.loads(content)
    except:
        return {"score": 0, "feedback": "评分失败"}

def review_code(question_text: str, reference_code: str, student_code: str) -> Dict[str, Any]:
    """
    AI 代码评审
    """
    from dotenv import load_dotenv
    env_path = Path(__file__).parent.parent / ".env"
    load_dotenv(dotenv_path=env_path, override=True)
    api_key = os.getenv("DEEPSEEK_API_KEY")
    
    prompt = f"""
请对学生提交的代码进行 Code Review。
题目：{question_text}
参考代码：
{reference_code}

学生代码：
{student_code}

请扮演一位资深技术面试官，对学生提交的代码进行 Code Review。
从以下维度评分（总分 0-10 分）：
1. 功能正确性 (40%)
2. 代码质量与规范 (30%)
3. 效率与优化 (20%)
4. 边界情况 (10%)

要求：
- 即使实现方式不同，只要逻辑正确且高效，也应给予高分。
- 评语应【简洁明了】，避免冗长。使用 Markdown 列表形式列出：
  - 【优点】：...
  - 【不足】：...
  - 【改进建议】：...

输出 JSON 格式：{{ "score": 8, "feedback": "..." }}
"""
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    try:
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=30)
        content = resp.json()["choices"][0]["message"]["content"]
        if content.startswith("```"): content = content.strip("`").replace("json", "")
        return json.loads(content)
    except:
        return {"score": 0, "feedback": "评审失败"}

def export_quiz_to_word(quiz_data: Dict[str, Any], output_path: str, include_answers: bool = True):
    """
    将题目数据导出为 Word 文档
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx not installed.")
        return False

    doc = Document()
    
    # 设置全文档默认字体为宋体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman' # 西文
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') # 中文
    
    # 标题
    title = quiz_data.get("title", "练习题")
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题也设置一下字体
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 描述
    if quiz_data.get("description"):
        p = doc.add_paragraph(quiz_data["description"])
    
    doc.add_paragraph("-" * 50)

    questions = quiz_data.get("questions", [])
    
    # 分类题目
    grouped_questions = {
        "单选题": [q for q in questions if q["type"] == "multiple_choice"],
        "多选题": [q for q in questions if q["type"] == "multi_select"],
        "判断题": [q for q in questions if q["type"] == "true_false"],
        "填空题": [q for q in questions if q["type"] == "fill_in_blank"],
        "简答题": [q for q in questions if q["type"] == "short_answer"],
        "代码题": [q for q in questions if q["type"] == "coding"],
    }
    
    chinese_numbers = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    section_index = 0

    def add_question_block(q, index, q_type_name):
        # 题干
        p = doc.add_paragraph()
        run = p.add_run(f"{index}. {q['stem']}")
        run.font.size = Pt(12)
        
        # 选项 (单选/多选)
        if q.get("options_json"):
            try:
                options = json.loads(q["options_json"])
                for i, opt in enumerate(options):
                    # A, B, C...
                    letter = chr(65 + i)
                    doc.add_paragraph(f"   {letter}. {opt}")
            except:
                pass
        
        # 答案和解析
        if include_answers:
            ans_p = doc.add_paragraph()
            
            # 格式化答案显示
            ans_text = q.get('answer')
            if q_type_name == "多选题":
                try:
                    # 如果是 JSON 数组字符串，转为 A,B,C
                    ans_list = json.loads(ans_text) if isinstance(ans_text, str) and ans_text.startswith("[") else ans_text
                    if isinstance(ans_list, list):
                        ans_text = ", ".join(ans_list)
                except:
                    pass
            elif q_type_name == "判断题":
                ans_text = "正确" if str(ans_text).lower() == "true" else "错误"

            ans_run = ans_p.add_run(f"【答案】 {ans_text}")
            ans_run.font.bold = True
            ans_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark Green
            
            if q.get("explanation"):
                exp_p = doc.add_paragraph()
                exp_run = exp_p.add_run(f"【解析】 {q['explanation']}")
                exp_run.font.italic = True
                exp_run.font.color.rgb = RGBColor(100, 100, 100) # Gray
        
        doc.add_paragraph() # 空行

    # 遍历所有类型并生成文档
    for type_name, q_list in grouped_questions.items():
        if q_list:
            section_title = f"{chinese_numbers[section_index]}、{type_name}"
            doc.add_heading(section_title, level=1)
            section_index += 1
            
            for i, q in enumerate(q_list, 1):
                add_question_block(q, i, type_name)

    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error saving word document: {e}")
        return False
