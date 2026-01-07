import os
import sys
import re
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed. Please run 'pip install python-docx'")
    sys.exit(1)

def set_chinese_font(run, font_name='宋体', ascii_font='Times New Roman'):
    """Helper to set Chinese font for a run"""
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_formatted_text(paragraph, text, default_color=None):
    """Parses **bold**, *italic*, and `code` text and adds runs to paragraph"""
    # Pattern to match ***bold italic***, **bold**, *italic*, and `code`
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    
    # Split the text by the pattern, but keep the delimiters
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
            set_chinese_font(run)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            set_chinese_font(run)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            set_chinese_font(run)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            # Use a mono font for code
            set_chinese_font(run, ascii_font='Consolas')
            run.font.color.rgb = RGBColor(199, 37, 78) # Dark pink/red for code
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run)
        
        if default_color and not (part.startswith('`') and part.endswith('`')):
            run.font.color.rgb = default_color

def md_to_word(md_file, word_file):
    doc = Document()
    
    # Global style settings for Chinese
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"Error reading file {md_file}: {e}")
        return

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                set_chinese_font(run, '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                set_chinese_font(run, '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                set_chinese_font(run, '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Checkboxes
        elif line.startswith('- [x]'):
            p = doc.add_paragraph()
            # Prefix
            run = p.add_run('☑ ')
            set_chinese_font(run)
            run.font.color.rgb = RGBColor(0, 128, 0) # Green for done
            
            # Content
            content = line[5:].strip()
            add_formatted_text(p, content, RGBColor(0, 128, 0))
            
        elif line.startswith('- [ ]'):
            p = doc.add_paragraph()
            # Prefix
            run = p.add_run('☐ ')
            set_chinese_font(run)
            
            # Content
            content = line[5:].strip()
            add_formatted_text(p, content)
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:].strip()
            add_formatted_text(p, content)
            
        # Normal Text
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    try:
        doc.save(word_file)
        print(f"Successfully converted: {md_file}")
        print(f"Saved to: {word_file}")
    except Exception as e:
        print(f"Error saving file {word_file}: {e}")

if __name__ == "__main__":
    # Default paths: Look for DEVELOPMENT_TASKS.md in the project root
    # script is in /scripts, so root is ../
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    
    default_md = os.path.join(project_root, "DEVELOPMENT_TASKS.md")
    default_docx = os.path.join(project_root, "DEVELOPMENT_TASKS.docx")
    
    md_path = default_md
    docx_path = default_docx
    
    # Allow command line args override
    if len(sys.argv) >= 2:
        md_path = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
        
    if not os.path.exists(md_path):
        print(f"Markdown file not found: {md_path}")
    else:
        md_to_word(md_path, docx_path)
